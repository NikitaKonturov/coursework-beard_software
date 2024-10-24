from docx import Document # Импорт класса Document в программу

def txt_to_docx(txtFile, docxFile): # Определение будущей функции
    doc = Document()
    
    with open(txtFile, 'r', encoding='utf-8') as file: # With гарантирует, что файл автоматически закроется
        # 'r' открывает файл в режиме для чтения, encoding указывает, что текст должен читаться именно в указанной кодировке
        for line in file:
            
            everyNewLine = line.strip()
            
            if everyNewLine == '':
                doc.add_paragraph()
            
            else:
                doc.add_paragraph(line.rstrip())
                    
    doc.save(docxFile)
    
txt_to_docx('test.txt', 'output.docx')